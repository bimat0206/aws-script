#!/usr/bin/env python3
"""
Comprehensive EC2 Cost and Configuration Report Generator
Scans all EC2 instances across all AWS regions, fetches detailed configurations including:
- Security Groups with detailed rules analysis
- AMI information and operating system detection  
- Instance-specific settings (user data, key pairs, IAM roles, networking)
- Performance metrics, costs, and generates security & optimization recommendations
"""

import boto3
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import sys
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

class ComprehensiveEC2CostReportGenerator:
    def __init__(self, profile_name='default', region='us-east-1', start_month=4, end_month=6, year=2025):
        self.profile_name = profile_name
        self.region = region
        self.all_regions = []
        
        try:
            # Create session with profile and region
            self.session = boto3.Session(profile_name=self.profile_name, region_name=self.region)
            
            # Test the session by getting account ID
            sts_client = self.session.client('sts')
            self.account_id = sts_client.get_caller_identity()['Account']
            
            # Initialize main region clients
            self.ec2 = self.session.client('ec2')
            self.ce = self.session.client('ce', region_name='us-east-1')  # Cost Explorer only in us-east-1
            self.cloudwatch = self.session.client('cloudwatch')
            
            # Get all available regions
            self.all_regions = self._get_all_regions()
            
            print(f"Successfully connected to AWS using profile '{self.profile_name}' in region '{self.region}'")
            print(f"Account ID: {self.account_id}")
            print(f"Will scan {len(self.all_regions)} regions for EC2 instances")
            
        except Exception as e:
            print(f"Error creating AWS session: {e}")
            print(f"Please ensure the AWS profile '{self.profile_name}' is configured correctly.")
            print(f"You can also try:")
            print(f"  1. Set a default region: aws configure set region {self.region} --profile {self.profile_name}")
            print(f"  2. Use a different region with --region parameter")
            print(f"  3. Check your AWS credentials: aws sts get-caller-identity --profile {self.profile_name}")
            sys.exit(1)
            
        self.start_date = datetime(year, start_month, 1)
        self.end_date = datetime(year, end_month + 1, 1)

    def _get_all_regions(self):
        """Get all available EC2 regions."""
        try:
            response = self.ec2.describe_regions()
            regions = [region['RegionName'] for region in response['Regions']]
            return sorted(regions)
        except Exception as e:
            print(f"Error fetching regions: {e}")
            # Fallback to common regions
            return ['us-east-1', 'us-west-1', 'us-west-2', 'eu-west-1', 'eu-central-1', 'ap-southeast-1', 'ap-northeast-1']

    def get_instance_performance_metrics(self, instance_id, region_name):
        """Get CloudWatch performance metrics for an instance."""
        try:
            cw_client = self.session.client('cloudwatch', region_name=region_name)
            end_time = datetime.utcnow()
            start_time = end_time - timedelta(days=7)
            
            metrics = {}
            
            # CPU Utilization
            cpu_response = cw_client.get_metric_statistics(
                Namespace='AWS/EC2',
                MetricName='CPUUtilization',
                Dimensions=[{'Name': 'InstanceId', 'Value': instance_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=3600,  # Hourly
                Statistics=['Average', 'Maximum']
            )
            
            if cpu_response['Datapoints']:
                cpu_data = cpu_response['Datapoints']
                metrics['avg_cpu'] = round(sum(dp['Average'] for dp in cpu_data) / len(cpu_data), 2)
                metrics['max_cpu'] = round(max(dp['Maximum'] for dp in cpu_data), 2)
            else:
                metrics['avg_cpu'] = 0
                metrics['max_cpu'] = 0
            
            # Network In/Out
            network_in = cw_client.get_metric_statistics(
                Namespace='AWS/EC2',
                MetricName='NetworkIn',
                Dimensions=[{'Name': 'InstanceId', 'Value': instance_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=3600,
                Statistics=['Sum']
            )
            
            network_out = cw_client.get_metric_statistics(
                Namespace='AWS/EC2',
                MetricName='NetworkOut',
                Dimensions=[{'Name': 'InstanceId', 'Value': instance_id}],
                StartTime=start_time,
                EndTime=end_time,
                Period=3600,
                Statistics=['Sum']
            )
            
            if network_in['Datapoints']:
                total_in = sum(dp['Sum'] for dp in network_in['Datapoints'])
                metrics['network_in_gb'] = round(total_in / (1024**3), 3)
            else:
                metrics['network_in_gb'] = 0
                
            if network_out['Datapoints']:
                total_out = sum(dp['Sum'] for dp in network_out['Datapoints'])
                metrics['network_out_gb'] = round(total_out / (1024**3), 3)
            else:
                metrics['network_out_gb'] = 0
            
            return metrics
        except Exception as e:
            print(f"    Could not get metrics for instance {instance_id}: {e}")
            return {
                'avg_cpu': 0,
                'max_cpu': 0,
                'network_in_gb': 0,
                'network_out_gb': 0
            }

    def get_instance_storage_details(self, instance, ec2_client):
        """Get detailed storage information for an instance."""
        storage_info = {
            'ebs_volumes': [],
            'total_ebs_size_gb': 0,
            'root_device_type': instance.get('RootDeviceType', 'Unknown'),
            'root_device_name': instance.get('RootDeviceName', 'Unknown')
        }
        
        try:
            # Get EBS volumes
            for bdm in instance.get('BlockDeviceMappings', []):
                if 'Ebs' in bdm:
                    volume_id = bdm['Ebs']['VolumeId']
                    try:
                        volume_response = ec2_client.describe_volumes(VolumeIds=[volume_id])
                        if volume_response['Volumes']:
                            volume = volume_response['Volumes'][0]
                            volume_info = {
                                'volume_id': volume_id,
                                'device_name': bdm['DeviceName'],
                                'size_gb': volume['Size'],
                                'volume_type': volume['VolumeType'],
                                'iops': volume.get('Iops', 'N/A'),
                                'encrypted': volume.get('Encrypted', False),
                                'state': volume['State']
                            }
                            storage_info['ebs_volumes'].append(volume_info)
                            storage_info['total_ebs_size_gb'] += volume['Size']
                    except Exception as e:
                        print(f"    Could not get volume details for {volume_id}: {e}")
        except Exception as e:
            print(f"    Could not get storage details: {e}")
        
        return storage_info

    def get_security_group_details(self, security_groups, ec2_client):
        """Get detailed security group information."""
        sg_details = []
        
        try:
            sg_ids = [sg['GroupId'] for sg in security_groups]
            if sg_ids:
                sg_response = ec2_client.describe_security_groups(GroupIds=sg_ids)
                
                for sg in sg_response['SecurityGroups']:
                    sg_info = {
                        'GroupId': sg['GroupId'],
                        'GroupName': sg['GroupName'],
                        'Description': sg['Description'],
                        'VpcId': sg.get('VpcId', 'EC2-Classic'),
                        'InboundRules': [],
                        'OutboundRules': []
                    }
                    
                    # Process inbound rules
                    for rule in sg.get('IpPermissions', []):
                        rule_info = {
                            'Protocol': rule.get('IpProtocol', 'Unknown'),
                            'FromPort': rule.get('FromPort', 'All'),
                            'ToPort': rule.get('ToPort', 'All'),
                            'Sources': []
                        }
                        
                        # IP ranges
                        for ip_range in rule.get('IpRanges', []):
                            rule_info['Sources'].append(f"CIDR: {ip_range.get('CidrIp', 'Unknown')}")
                        
                        # IPv6 ranges
                        for ipv6_range in rule.get('Ipv6Ranges', []):
                            rule_info['Sources'].append(f"IPv6: {ipv6_range.get('CidrIpv6', 'Unknown')}")
                        
                        # Security group references
                        for sg_ref in rule.get('UserIdGroupPairs', []):
                            rule_info['Sources'].append(f"SG: {sg_ref.get('GroupId', 'Unknown')}")
                        
                        # Prefix lists
                        for prefix in rule.get('PrefixListIds', []):
                            rule_info['Sources'].append(f"PL: {prefix.get('PrefixListId', 'Unknown')}")
                        
                        sg_info['InboundRules'].append(rule_info)
                    
                    # Process outbound rules
                    for rule in sg.get('IpPermissionsEgress', []):
                        rule_info = {
                            'Protocol': rule.get('IpProtocol', 'Unknown'),
                            'FromPort': rule.get('FromPort', 'All'),
                            'ToPort': rule.get('ToPort', 'All'),
                            'Destinations': []
                        }
                        
                        # IP ranges
                        for ip_range in rule.get('IpRanges', []):
                            rule_info['Destinations'].append(f"CIDR: {ip_range.get('CidrIp', 'Unknown')}")
                        
                        # IPv6 ranges
                        for ipv6_range in rule.get('Ipv6Ranges', []):
                            rule_info['Destinations'].append(f"IPv6: {ipv6_range.get('CidrIpv6', 'Unknown')}")
                        
                        # Security group references
                        for sg_ref in rule.get('UserIdGroupPairs', []):
                            rule_info['Destinations'].append(f"SG: {sg_ref.get('GroupId', 'Unknown')}")
                        
                        # Prefix lists
                        for prefix in rule.get('PrefixListIds', []):
                            rule_info['Destinations'].append(f"PL: {prefix.get('PrefixListId', 'Unknown')}")
                        
                        sg_info['OutboundRules'].append(rule_info)
                    
                    sg_details.append(sg_info)
        
        except Exception as e:
            print(f"    Could not get security group details: {e}")
        
        return sg_details

    def get_ami_detailed_info(self, ami_id, ec2_client):
        """Get comprehensive AMI information."""
        ami_info = {
            'ImageId': ami_id,
            'Name': 'Unknown',
            'Description': 'Unknown',
            'Owner': 'Unknown',
            'Public': False,
            'Platform': 'Unknown',
            'Architecture': 'Unknown',
            'VirtualizationType': 'Unknown',
            'Hypervisor': 'Unknown',
            'RootDeviceType': 'Unknown',
            'RootDeviceName': 'Unknown',
            'CreationDate': 'Unknown',
            'State': 'Unknown',
            'BlockDeviceMappings': [],
            'Tags': {}
        }
        
        try:
            ami_response = ec2_client.describe_images(ImageIds=[ami_id])
            if ami_response['Images']:
                ami = ami_response['Images'][0]
                ami_info.update({
                    'Name': ami.get('Name', 'Unknown'),
                    'Description': ami.get('Description', 'Unknown')[:200],  # Truncate long descriptions
                    'Owner': ami.get('OwnerId', 'Unknown'),
                    'Public': ami.get('Public', False),
                    'Platform': ami.get('Platform', ami.get('PlatformDetails', 'Linux/UNIX')),
                    'Architecture': ami.get('Architecture', 'Unknown'),
                    'VirtualizationType': ami.get('VirtualizationType', 'Unknown'),
                    'Hypervisor': ami.get('Hypervisor', 'Unknown'),
                    'RootDeviceType': ami.get('RootDeviceType', 'Unknown'),
                    'RootDeviceName': ami.get('RootDeviceName', 'Unknown'),
                    'CreationDate': ami.get('CreationDate', 'Unknown'),
                    'State': ami.get('State', 'Unknown'),
                    'BlockDeviceMappings': ami.get('BlockDeviceMappings', [])
                })
                
                # Get AMI tags
                ami_tags = {tag['Key']: tag['Value'] for tag in ami.get('Tags', [])}
                ami_info['Tags'] = ami_tags
        
        except Exception as e:
            print(f"    Could not get AMI details for {ami_id}: {e}")
        
        return ami_info

    def detect_operating_system(self, ami_info, platform):
        """Detect operating system from AMI and platform information."""
        ami_name = ami_info.get('Name', '').lower()
        ami_description = ami_info.get('Description', '').lower()
        platform_details = ami_info.get('Platform', '').lower()
        
        # Windows detection
        if platform and 'windows' in platform.lower():
            if 'server 2022' in ami_name or 'server 2022' in ami_description:
                return 'Windows Server 2022'
            elif 'server 2019' in ami_name or 'server 2019' in ami_description:
                return 'Windows Server 2019'
            elif 'server 2016' in ami_name or 'server 2016' in ami_description:
                return 'Windows Server 2016'
            elif 'server 2012' in ami_name or 'server 2012' in ami_description:
                return 'Windows Server 2012'
            elif 'windows 11' in ami_name or 'windows 11' in ami_description:
                return 'Windows 11'
            elif 'windows 10' in ami_name or 'windows 10' in ami_description:
                return 'Windows 10'
            else:
                return 'Windows (Unknown Version)'
        
        # Linux distributions
        elif 'amazon' in ami_name or 'amzn' in ami_name:
            if 'amazon linux 2023' in ami_name or 'al2023' in ami_name:
                return 'Amazon Linux 2023'
            elif 'amazon linux 2' in ami_name or 'amzn2' in ami_name:
                return 'Amazon Linux 2'
            else:
                return 'Amazon Linux'
        elif 'ubuntu' in ami_name:
            if '22.04' in ami_name or 'jammy' in ami_name:
                return 'Ubuntu 22.04 LTS (Jammy)'
            elif '20.04' in ami_name or 'focal' in ami_name:
                return 'Ubuntu 20.04 LTS (Focal)'
            elif '18.04' in ami_name or 'bionic' in ami_name:
                return 'Ubuntu 18.04 LTS (Bionic)'
            else:
                return 'Ubuntu (Unknown Version)'
        elif 'rhel' in ami_name or 'red hat' in ami_name:
            if 'rhel-9' in ami_name or 'rhel 9' in ami_name:
                return 'Red Hat Enterprise Linux 9'
            elif 'rhel-8' in ami_name or 'rhel 8' in ami_name:
                return 'Red Hat Enterprise Linux 8'
            elif 'rhel-7' in ami_name or 'rhel 7' in ami_name:
                return 'Red Hat Enterprise Linux 7'
            else:
                return 'Red Hat Enterprise Linux'
        elif 'centos' in ami_name:
            if 'centos 8' in ami_name:
                return 'CentOS 8'
            elif 'centos 7' in ami_name:
                return 'CentOS 7'
            else:
                return 'CentOS'
        elif 'debian' in ami_name:
            if 'debian-11' in ami_name or 'bullseye' in ami_name:
                return 'Debian 11 (Bullseye)'
            elif 'debian-10' in ami_name or 'buster' in ami_name:
                return 'Debian 10 (Buster)'
            else:
                return 'Debian'
        elif 'suse' in ami_name or 'sles' in ami_name:
            return 'SUSE Linux Enterprise Server'
        elif 'freebsd' in ami_name:
            return 'FreeBSD'
        elif 'oracle' in ami_name:
            return 'Oracle Linux'
        else:
            return 'Linux/Unix (Unknown Distribution)'

    def get_instance_specific_settings(self, instance_id, ec2_client):
        """Get instance-specific settings like user data, metadata options, etc."""
        settings = {
            'UserData': 'None',
            'UserDataSize': 0,
            'SourceDestCheck': 'Unknown',
            'EbsOptimized': 'Unknown',
            'SriovNetSupport': 'Unknown',
            'EnaSupport': 'Unknown',
            'MetadataOptions': {},
            'MaintenanceOptions': {},
            'BootMode': 'Unknown'
        }
        
        try:
            # Get user data
            try:
                user_data_response = ec2_client.describe_instance_attribute(
                    InstanceId=instance_id,
                    Attribute='userData'
                )
                user_data = user_data_response.get('UserData', {}).get('Value')
                if user_data:
                    settings['UserData'] = 'Present'
                    settings['UserDataSize'] = len(user_data)
                else:
                    settings['UserData'] = 'None'
            except:
                settings['UserData'] = 'Unknown'
            
            # Get source/destination check
            try:
                src_dest_response = ec2_client.describe_instance_attribute(
                    InstanceId=instance_id,
                    Attribute='sourceDestCheck'
                )
                settings['SourceDestCheck'] = src_dest_response.get('SourceDestCheck', {}).get('Value', 'Unknown')
            except:
                pass
            
            # Get EBS optimization
            try:
                ebs_opt_response = ec2_client.describe_instance_attribute(
                    InstanceId=instance_id,
                    Attribute='ebsOptimized'
                )
                settings['EbsOptimized'] = ebs_opt_response.get('EbsOptimized', {}).get('Value', 'Unknown')
            except:
                pass
            
            # Get SRIOV support
            try:
                sriov_response = ec2_client.describe_instance_attribute(
                    InstanceId=instance_id,
                    Attribute='sriovNetSupport'
                )
                settings['SriovNetSupport'] = sriov_response.get('SriovNetSupport', {}).get('Value', 'simple')
            except:
                pass
            
            # Get ENA support
            try:
                ena_response = ec2_client.describe_instance_attribute(
                    InstanceId=instance_id,
                    Attribute='enaSupport'
                )
                settings['EnaSupport'] = ena_response.get('EnaSupport', {}).get('Value', 'Unknown')
            except:
                pass
            
        except Exception as e:
            print(f"    Could not get instance-specific settings: {e}")
        
        return settings

    def get_instance_detailed_config(self, instance, region_name):
        """Get comprehensive instance configuration."""
        config = {}
        
        try:
            ec2_client = self.session.client('ec2', region_name=region_name)
            instance_id = instance['InstanceId']
            
            # Basic instance information
            config['InstanceId'] = instance_id
            config['Region'] = region_name
            config['InstanceType'] = instance['InstanceType']
            config['State'] = instance['State']['Name']
            config['LaunchTime'] = instance['LaunchTime'].strftime("%Y-%m-%d %H:%M:%S")
            config['ImageId'] = instance['ImageId']
            config['KeyName'] = instance.get('KeyName', 'None')
            config['Platform'] = instance.get('Platform', 'Linux/Unix')
            config['Architecture'] = instance.get('Architecture', 'Unknown')
            config['VirtualizationType'] = instance.get('VirtualizationType', 'Unknown')
            config['Hypervisor'] = instance.get('Hypervisor', 'Unknown')
            
            # Calculate uptime
            if config['State'] == 'running':
                uptime = datetime.now(instance['LaunchTime'].tzinfo) - instance['LaunchTime']
                config['UptimeDays'] = uptime.days
            else:
                config['UptimeDays'] = 0
            
            # Network configuration
            config['VpcId'] = instance.get('VpcId', 'EC2-Classic')
            config['SubnetId'] = instance.get('SubnetId', 'None')
            config['AvailabilityZone'] = instance.get('Placement', {}).get('AvailabilityZone', 'Unknown')
            config['Tenancy'] = instance.get('Placement', {}).get('Tenancy', 'default')
            config['PlacementGroup'] = instance.get('Placement', {}).get('GroupName', 'None')
            config['PlacementPartitionNumber'] = instance.get('Placement', {}).get('PartitionNumber', 'None')
            config['HostId'] = instance.get('Placement', {}).get('HostId', 'None')
            
            # IP addresses and DNS
            config['PrivateIpAddress'] = instance.get('PrivateIpAddress', 'None')
            config['PublicIpAddress'] = instance.get('PublicIpAddress', 'None')
            config['PrivateDnsName'] = instance.get('PrivateDnsName', 'None')
            config['PublicDnsName'] = instance.get('PublicDnsName', 'None')
            
            # Network interfaces
            network_interfaces = instance.get('NetworkInterfaces', [])
            config['NetworkInterfaceCount'] = len(network_interfaces)
            config['PrimaryNetworkInterface'] = network_interfaces[0].get('NetworkInterfaceId', 'None') if network_interfaces else 'None'
            
            # Security groups with detailed information
            security_groups = instance.get('SecurityGroups', [])
            config['SecurityGroups'] = ', '.join([sg['GroupName'] for sg in security_groups]) if security_groups else 'None'
            config['SecurityGroupCount'] = len(security_groups)
            config['SecurityGroupDetails'] = self.get_security_group_details(security_groups, ec2_client)
            
            # IAM role detailed information
            iam_info = instance.get('IamInstanceProfile', {})
            config['IamRole'] = iam_info.get('Arn', 'None').split('/')[-1] if iam_info.get('Arn') else 'None'
            config['IamInstanceProfileArn'] = iam_info.get('Arn', 'None')
            
            # Monitoring and metadata
            config['MonitoringState'] = instance.get('Monitoring', {}).get('State', 'disabled')
            config['MetadataOptions'] = instance.get('MetadataOptions', {})
            config['MaintenanceOptions'] = instance.get('MaintenanceOptions', {})
            config['BootMode'] = instance.get('BootMode', 'legacy-bios')
            
            # Storage details
            storage_info = self.get_instance_storage_details(instance, ec2_client)
            config.update(storage_info)
            
            # Enhanced Tags
            tags = {tag['Key']: tag['Value'] for tag in instance.get('Tags', [])}
            config['Tags'] = tags
            config['TagCount'] = len(tags)
            config['Name'] = tags.get('Name', 'Unnamed')
            config['Environment'] = tags.get('Environment', tags.get('Env', 'Untagged'))
            config['Owner'] = tags.get('Owner', 'Untagged')
            config['Project'] = tags.get('Project', 'Untagged')
            config['CostCenter'] = tags.get('CostCenter', tags.get('Cost-Center', 'Untagged'))
            config['Application'] = tags.get('Application', tags.get('App', 'Untagged'))
            
            # AMI detailed information
            config['AmiDetails'] = self.get_ami_detailed_info(instance['ImageId'], ec2_client)
            
            # Operating System detection
            config['OperatingSystem'] = self.detect_operating_system(config['AmiDetails'], config['Platform'])
            
            # Instance-specific settings
            instance_settings = self.get_instance_specific_settings(instance_id, ec2_client)
            config.update(instance_settings)
            
            # Performance metrics
            if config['State'] == 'running':
                performance_metrics = self.get_instance_performance_metrics(instance_id, region_name)
                config.update(performance_metrics)
            else:
                config.update({
                    'avg_cpu': 0,
                    'max_cpu': 0,
                    'network_in_gb': 0,
                    'network_out_gb': 0
                })
            
            # Additional instance attributes
            try:
                shutdown_response = ec2_client.describe_instance_attribute(
                    InstanceId=instance_id,
                    Attribute='instanceInitiatedShutdownBehavior'
                )
                config['ShutdownBehavior'] = shutdown_response.get('InstanceInitiatedShutdownBehavior', {}).get('Value', 'stop')
            except:
                config['ShutdownBehavior'] = 'Unknown'
                
        except Exception as e:
            print(f"    Error getting instance details: {e}")
            
        return config

    def scan_region_instances(self, region_name):
        """Scan all EC2 instances in a specific region."""
        instances_data = []
        
        try:
            print(f"  Scanning region: {region_name}")
            ec2_client = self.session.client('ec2', region_name=region_name)
            
            # Get all instances in the region
            paginator = ec2_client.get_paginator('describe_instances')
            page_iterator = paginator.paginate()
            
            instance_count = 0
            for page in page_iterator:
                for reservation in page['Reservations']:
                    for instance in reservation['Instances']:
                        instance_count += 1
                        print(f"    Processing instance {instance_count}: {instance['InstanceId']} in {region_name}")
                        
                        try:
                            config = self.get_instance_detailed_config(instance, region_name)
                            instances_data.append(config)
                        except Exception as e:
                            print(f"    Error processing instance {instance['InstanceId']}: {e}")
                            instances_data.append({
                                'InstanceId': instance['InstanceId'],
                                'Region': region_name,
                                'Error': str(e)
                            })
            
            if instance_count == 0:
                print(f"    No instances found in {region_name}")
                
        except Exception as e:
            print(f"    Error scanning region {region_name}: {e}")
        
        return instances_data

    def fetch_all_ec2_instances(self):
        """Fetch all EC2 instances across all regions with parallel processing."""
        print("Scanning all EC2 instances across all AWS regions...")
        all_instances = []
        
        # Use ThreadPoolExecutor for parallel region scanning
        with ThreadPoolExecutor(max_workers=5) as executor:
            future_to_region = {executor.submit(self.scan_region_instances, region): region for region in self.all_regions}
            
            for future in as_completed(future_to_region):
                region = future_to_region[future]
                try:
                    region_instances = future.result()
                    all_instances.extend(region_instances)
                except Exception as e:
                    print(f"Error scanning region {region}: {e}")
        
        print(f"Found {len(all_instances)} total EC2 instances across all regions")
        return all_instances

    def get_ec2_cost_history(self):
        """Gets EC2 cost data for the specified period."""
        print(f"Fetching EC2 cost data from {self.start_date.strftime('%Y-%m')} to {self.end_date.strftime('%Y-%m')}...")
        try:
            response = self.ce.get_cost_and_usage(
                TimePeriod={
                    'Start': self.start_date.strftime('%Y-%m-%d'),
                    'End': self.end_date.strftime('%Y-%m-%d')
                },
                Granularity='MONTHLY',
                Filter={
                    'Dimensions': {
                        'Key': 'SERVICE',
                        'Values': ['Amazon Elastic Compute Cloud - Compute']
                    }
                },
                Metrics=['UnblendedCost'],
                GroupBy=[{'Type': 'DIMENSION', 'Key': 'USAGE_TYPE'}]
            )
            return response['ResultsByTime']
        except Exception as e:
            print(f"Error fetching cost data: {e}")
            return None

    def format_size(self, size_gb):
        """Format storage size in human readable format."""
        if size_gb == 0:
            return "0 GB"
        elif size_gb < 1024:
            return f"{size_gb} GB"
        else:
            return f"{size_gb/1024:.2f} TB"

    def generate_report(self, instances_data, cost_data):
        """Generates a comprehensive Word document report."""
        print("Generating comprehensive EC2 report...")
        doc = Document()
        doc.add_heading('AWS EC2 Comprehensive Configuration and Cost Report', 0)

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"AWS Account ID: {self.account_id}")
        doc.add_paragraph(f"Profile: {self.profile_name}")
        doc.add_paragraph(f"Reporting Period: {self.start_date.strftime('%B %Y')} - {(self.end_date - timedelta(days=1)).strftime('%B %Y')}")
        doc.add_paragraph(f"Regions Scanned: {len(self.all_regions)}")
        
        if instances_data:
            total_instances = len(instances_data)
            running_instances = sum(1 for inst in instances_data if inst.get('State') == 'running')
            stopped_instances = sum(1 for inst in instances_data if inst.get('State') == 'stopped')
            total_storage = sum(inst.get('total_ebs_size_gb', 0) for inst in instances_data)
            regions_with_instances = len(set(inst.get('Region') for inst in instances_data if inst.get('Region')))
            
            # Instance type distribution
            instance_types = {}
            for inst in instances_data:
                inst_type = inst.get('InstanceType', 'Unknown')
                instance_types[inst_type] = instance_types.get(inst_type, 0) + 1
            
            doc.add_paragraph(f"Total EC2 Instances: {total_instances}")
            doc.add_paragraph(f"Running Instances: {running_instances}")
            doc.add_paragraph(f"Stopped Instances: {stopped_instances}")
            doc.add_paragraph(f"Regions with Instances: {regions_with_instances}")
            doc.add_paragraph(f"Total EBS Storage: {self.format_size(total_storage)}")
            doc.add_paragraph(f"Most Common Instance Type: {max(instance_types.items(), key=lambda x: x[1])[0] if instance_types else 'None'}")

        # Regional Distribution
        doc.add_heading('Regional Distribution', level=1)
        if instances_data:
            region_stats = {}
            for inst in instances_data:
                region = inst.get('Region', 'Unknown')
                if region not in region_stats:
                    region_stats[region] = {'total': 0, 'running': 0, 'stopped': 0}
                region_stats[region]['total'] += 1
                if inst.get('State') == 'running':
                    region_stats[region]['running'] += 1
                elif inst.get('State') == 'stopped':
                    region_stats[region]['stopped'] += 1
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ['Region', 'Total Instances', 'Running', 'Stopped']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for region, stats in sorted(region_stats.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = region
                row_cells[1].text = str(stats['total'])
                row_cells[2].text = str(stats['running'])
                row_cells[3].text = str(stats['stopped'])

        # Instance Overview Table
        doc.add_heading('EC2 Instance Overview', level=1)
        if instances_data:
            table = doc.add_table(rows=1, cols=12)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ['Name', 'Instance ID', 'Type', 'State', 'OS', 'Region', 'AZ', 'Storage (GB)', 'Avg CPU %', 'Key Pair', 'Launch Time', 'Uptime (Days)']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Sort by region, then by name
            sorted_instances = sorted(instances_data, key=lambda x: (x.get('Region', ''), x.get('Name', x.get('InstanceId', ''))))
            
            for instance in sorted_instances:
                row_cells = table.add_row().cells
                row_cells[0].text = instance.get('Name', 'Unnamed')[:20]  # Truncate long names
                row_cells[1].text = instance.get('InstanceId', 'Unknown')
                row_cells[2].text = instance.get('InstanceType', 'Unknown')
                row_cells[3].text = instance.get('State', 'Unknown')
                row_cells[4].text = instance.get('OperatingSystem', 'Unknown')[:25]  # Truncate long OS names
                row_cells[5].text = instance.get('Region', 'Unknown')
                row_cells[6].text = instance.get('AvailabilityZone', 'Unknown')
                row_cells[7].text = str(instance.get('total_ebs_size_gb', 0))
                row_cells[8].text = f"{instance.get('avg_cpu', 0):.1f}"
                row_cells[9].text = instance.get('KeyName', 'None')
                row_cells[10].text = instance.get('LaunchTime', 'Unknown')[:10]  # Date only
                row_cells[11].text = str(instance.get('UptimeDays', 0))

        # Security Groups Analysis
        doc.add_heading('Security Groups Analysis', level=1)
        if instances_data:
            # Collect all unique security groups
            all_security_groups = {}
            for instance in instances_data:
                for sg_detail in instance.get('SecurityGroupDetails', []):
                    sg_id = sg_detail['GroupId']
                    if sg_id not in all_security_groups:
                        all_security_groups[sg_id] = sg_detail
                        # Add instance references
                        all_security_groups[sg_id]['UsedByInstances'] = []
                    all_security_groups[sg_id]['UsedByInstances'].append({
                        'InstanceId': instance.get('InstanceId'),
                        'Name': instance.get('Name', 'Unnamed'),
                        'Region': instance.get('Region')
                    })
            
            doc.add_paragraph(f"Total unique security groups across all instances: {len(all_security_groups)}")
            
            for sg_id, sg_info in sorted(all_security_groups.items()):
                doc.add_heading(f"Security Group: {sg_info['GroupName']} ({sg_id})", level=2)
                
                # Basic SG info
                doc.add_paragraph(f"Description: {sg_info['Description']}")
                doc.add_paragraph(f"VPC: {sg_info['VpcId']}")
                doc.add_paragraph(f"Used by {len(sg_info['UsedByInstances'])} instance(s)")
                
                # Inbound rules
                if sg_info['InboundRules']:
                    doc.add_heading('Inbound Rules', level=3)
                    inbound_table = doc.add_table(rows=1, cols=4)
                    inbound_table.style = 'Table Grid'
                    hdr_cells = inbound_table.rows[0].cells
                    hdr_cells[0].text = 'Protocol'
                    hdr_cells[1].text = 'Port Range'
                    hdr_cells[2].text = 'Source'
                    hdr_cells[3].text = 'Description'
                    for cell in hdr_cells:
                        cell.paragraphs[0].runs[0].font.bold = True
                    
                    for rule in sg_info['InboundRules']:
                        row_cells = inbound_table.add_row().cells
                        row_cells[0].text = rule['Protocol']
                        
                        if rule['FromPort'] == 'All':
                            row_cells[1].text = 'All'
                        elif rule['FromPort'] == rule['ToPort']:
                            row_cells[1].text = str(rule['FromPort'])
                        else:
                            row_cells[1].text = f"{rule['FromPort']}-{rule['ToPort']}"
                        
                        row_cells[2].text = '; '.join(rule['Sources'][:3])  # Limit to 3 sources for readability
                        row_cells[3].text = ''  # Could add rule descriptions if available
                
                # Outbound rules
                if sg_info['OutboundRules']:
                    doc.add_heading('Outbound Rules', level=3)
                    outbound_table = doc.add_table(rows=1, cols=4)
                    outbound_table.style = 'Table Grid'
                    hdr_cells = outbound_table.rows[0].cells
                    hdr_cells[0].text = 'Protocol'
                    hdr_cells[1].text = 'Port Range'
                    hdr_cells[2].text = 'Destination'
                    hdr_cells[3].text = 'Description'
                    for cell in hdr_cells:
                        cell.paragraphs[0].runs[0].font.bold = True
                    
                    for rule in sg_info['OutboundRules']:
                        row_cells = outbound_table.add_row().cells
                        row_cells[0].text = rule['Protocol']
                        
                        if rule['FromPort'] == 'All':
                            row_cells[1].text = 'All'
                        elif rule['FromPort'] == rule['ToPort']:
                            row_cells[1].text = str(rule['FromPort'])
                        else:
                            row_cells[1].text = f"{rule['FromPort']}-{rule['ToPort']}"
                        
                        row_cells[2].text = '; '.join(rule['Destinations'][:3])  # Limit for readability
                        row_cells[3].text = ''

        # Operating Systems Summary
        doc.add_heading('Operating Systems Summary', level=1)
        if instances_data:
            os_stats = {}
            for instance in instances_data:
                os = instance.get('OperatingSystem', 'Unknown')
                if os not in os_stats:
                    os_stats[os] = {'total': 0, 'running': 0, 'stopped': 0}
                os_stats[os]['total'] += 1
                if instance.get('State') == 'running':
                    os_stats[os]['running'] += 1
                elif instance.get('State') == 'stopped':
                    os_stats[os]['stopped'] += 1
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ['Operating System', 'Total Instances', 'Running', 'Stopped']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for os, stats in sorted(os_stats.items(), key=lambda x: x[1]['total'], reverse=True):
                row_cells = table.add_row().cells
                row_cells[0].text = os
                row_cells[1].text = str(stats['total'])
                row_cells[2].text = str(stats['running'])
                row_cells[3].text = str(stats['stopped'])

        # AMI Analysis
        doc.add_heading('AMI Analysis', level=1)
        if instances_data:
            ami_stats = {}
            for instance in instances_data:
                ami_details = instance.get('AmiDetails', {})
                ami_id = ami_details.get('ImageId', 'Unknown')
                if ami_id not in ami_stats:
                    ami_stats[ami_id] = {
                        'name': ami_details.get('Name', 'Unknown'),
                        'owner': ami_details.get('Owner', 'Unknown'),
                        'platform': ami_details.get('Platform', 'Unknown'),
                        'public': ami_details.get('Public', False),
                        'instances': []
                    }
                ami_stats[ami_id]['instances'].append({
                    'InstanceId': instance.get('InstanceId'),
                    'Name': instance.get('Name', 'Unnamed'),
                    'Region': instance.get('Region')
                })
            
            doc.add_paragraph(f"Total unique AMIs in use: {len(ami_stats)}")
            
            # Top 10 most used AMIs
            top_amis = sorted(ami_stats.items(), key=lambda x: len(x[1]['instances']), reverse=True)[:10]
            
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ['AMI ID', 'AMI Name', 'Owner', 'Platform', 'Public', 'Instance Count']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for ami_id, ami_info in top_amis:
                row_cells = table.add_row().cells
                row_cells[0].text = ami_id
                row_cells[1].text = ami_info['name'][:30]  # Truncate long names
                row_cells[2].text = ami_info['owner']
                row_cells[3].text = ami_info['platform']
                row_cells[4].text = 'Yes' if ami_info['public'] else 'No'
                row_cells[5].text = str(len(ami_info['instances']))

        # Detailed Instance Configurations (Enhanced)
        doc.add_heading('Detailed Instance Configurations (Sample)', level=1)
        sample_instances = sorted(instances_data, key=lambda x: x.get('avg_cpu', 0), reverse=True)[:10]
        
        for instance in sample_instances:
            doc.add_heading(f"Instance: {instance.get('Name', instance.get('InstanceId', 'Unknown'))}", level=2)
            
            # Basic Information
            doc.add_heading('Basic Information', level=3)
            basic_info = [
                f"Instance ID: {instance.get('InstanceId', 'Unknown')}",
                f"Instance Type: {instance.get('InstanceType', 'Unknown')}",
                f"State: {instance.get('State', 'Unknown')}",
                f"Region: {instance.get('Region', 'Unknown')}",
                f"Availability Zone: {instance.get('AvailabilityZone', 'Unknown')}",
                f"Launch Time: {instance.get('LaunchTime', 'Unknown')}",
                f"Uptime: {instance.get('UptimeDays', 0)} days",
                f"Operating System: {instance.get('OperatingSystem', 'Unknown')}",
                f"Architecture: {instance.get('Architecture', 'Unknown')}",
                f"Virtualization Type: {instance.get('VirtualizationType', 'Unknown')}",
            ]
            for info in basic_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # AMI Information
            doc.add_heading('AMI Information', level=3)
            ami_details = instance.get('AmiDetails', {})
            ami_info = [
                f"AMI ID: {ami_details.get('ImageId', 'Unknown')}",
                f"AMI Name: {ami_details.get('Name', 'Unknown')}",
                f"AMI Owner: {ami_details.get('Owner', 'Unknown')}",
                f"Platform: {ami_details.get('Platform', 'Unknown')}",
                f"Architecture: {ami_details.get('Architecture', 'Unknown')}",
                f"Root Device Type: {ami_details.get('RootDeviceType', 'Unknown')}",
                f"Public AMI: {'Yes' if ami_details.get('Public', False) else 'No'}",
                f"Creation Date: {ami_details.get('CreationDate', 'Unknown')}",
            ]
            for info in ami_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # Network Configuration
            doc.add_heading('Network Configuration', level=3)
            network_info = [
                f"VPC ID: {instance.get('VpcId', 'Unknown')}",
                f"Subnet ID: {instance.get('SubnetId', 'Unknown')}",
                f"Private IP: {instance.get('PrivateIpAddress', 'None')}",
                f"Public IP: {instance.get('PublicIpAddress', 'None')}",
                f"Private DNS: {instance.get('PrivateDnsName', 'None')}",
                f"Public DNS: {instance.get('PublicDnsName', 'None')}",
                f"Security Groups: {instance.get('SecurityGroups', 'None')}",
                f"Network Interfaces: {instance.get('NetworkInterfaceCount', 0)}",
            ]
            for info in network_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # Security Configuration
            doc.add_heading('Security & IAM Configuration', level=3)
            security_info = [
                f"Key Pair: {instance.get('KeyName', 'None')}",
                f"IAM Role: {instance.get('IamRole', 'None')}",
                f"IAM Instance Profile: {instance.get('IamInstanceProfileArn', 'None')}",
                f"Source/Dest Check: {instance.get('SourceDestCheck', 'Unknown')}",
                f"Security Group Count: {instance.get('SecurityGroupCount', 0)}",
            ]
            for info in security_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # Instance-Specific Settings
            doc.add_heading('Instance-Specific Settings', level=3)
            settings_info = [
                f"User Data: {instance.get('UserData', 'Unknown')} ({instance.get('UserDataSize', 0)} bytes)",
                f"EBS Optimized: {instance.get('EbsOptimized', 'Unknown')}",
                f"Enhanced Networking (ENA): {instance.get('EnaSupport', 'Unknown')}",
                f"SR-IOV Support: {instance.get('SriovNetSupport', 'Unknown')}",
                f"Monitoring: {instance.get('MonitoringState', 'Unknown')}",
                f"Shutdown Behavior: {instance.get('ShutdownBehavior', 'Unknown')}",
                f"Boot Mode: {instance.get('BootMode', 'Unknown')}",
                f"Tenancy: {instance.get('Tenancy', 'Unknown')}",
                f"Placement Group: {instance.get('PlacementGroup', 'None')}",
            ]
            for info in settings_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # Tags Information
            doc.add_heading('Tags', level=3)
            tags = instance.get('Tags', {})
            if tags:
                for key, value in sorted(tags.items()):
                    doc.add_paragraph(f"{key}: {value}", style='List Bullet')
            else:
                doc.add_paragraph("No tags assigned", style='List Bullet')
            
            # Performance Metrics
            if instance.get('State') == 'running':
                doc.add_heading('Performance Metrics (Last 7 Days)', level=3)
                perf_info = [
                    f"Average CPU Utilization: {instance.get('avg_cpu', 0):.2f}%",
                    f"Maximum CPU Utilization: {instance.get('max_cpu', 0):.2f}%",
                    f"Network In: {instance.get('network_in_gb', 0):.3f} GB",
                    f"Network Out: {instance.get('network_out_gb', 0):.3f} GB",
                ]
                for info in perf_info:
                    doc.add_paragraph(info, style='List Bullet')
            
            # Storage Information
            doc.add_heading('Storage Configuration', level=3)
            storage_info = [
                f"Root Device Type: {instance.get('root_device_type', 'Unknown')}",
                f"Root Device Name: {instance.get('root_device_name', 'Unknown')}",
                f"Total EBS Storage: {instance.get('total_ebs_size_gb', 0)} GB",
                f"EBS Volumes Count: {len(instance.get('ebs_volumes', []))}",
            ]
            for info in storage_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # EBS Volume Details
            ebs_volumes = instance.get('ebs_volumes', [])
            if ebs_volumes:
                doc.add_heading('EBS Volume Details', level=3)
                for volume in ebs_volumes:
                    vol_info = (f"Volume {volume.get('volume_id', 'Unknown')}: "
                              f"{volume.get('size_gb', 0)} GB {volume.get('volume_type', 'Unknown')} "
                              f"on {volume.get('device_name', 'Unknown')} "
                              f"({'Encrypted' if volume.get('encrypted', False) else 'Unencrypted'})")
                    doc.add_paragraph(vol_info, style='List Bullet')

        # Cost Analysis
        doc.add_heading(f"EC2 Cost Analysis ({self.start_date.strftime('%B')} - {(self.end_date - timedelta(days=1)).strftime('%B %Y')})", level=1)
        if cost_data:
            total_cost = 0
            for month_data in cost_data:
                month = datetime.strptime(month_data['TimePeriod']['Start'], '%Y-%m-%d').strftime('%B %Y')
                month_total = 0
                doc.add_heading(f"Costs for {month}", level=2)
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Usage Type'
                hdr_cells[1].text = 'Cost (USD)'
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True

                groups = sorted(month_data['Groups'], key=lambda x: float(x['Metrics']['UnblendedCost']['Amount']), reverse=True)

                for group in groups:
                    cost = float(group['Metrics']['UnblendedCost']['Amount'])
                    if cost > 0:
                        row_cells = table.add_row().cells
                        row_cells[0].text = group['Keys'][0]
                        row_cells[1].text = f"${cost:,.2f}"
                        month_total += cost
                
                total_cost += month_total
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.add_run(f'Total cost for {month}: ').bold = True
                p.add_run(f"${month_total:,.2f}")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run('Total EC2 cost for the period: ').bold = True
            p.add_run(f"${total_cost:,.2f}")
        else:
            doc.add_paragraph("Could not retrieve EC2 cost data.")

        # Recommendations
        doc.add_heading('Security & Optimization Recommendations', level=1)
        if instances_data:
            doc.add_paragraph("Based on the comprehensive analysis, here are security and optimization recommendations:")
            
            # Performance optimization
            underutilized = [inst for inst in instances_data 
                           if inst.get('State') == 'running' and inst.get('avg_cpu', 0) < 5]
            stopped_instances = [inst for inst in instances_data if inst.get('State') == 'stopped']
            
            if underutilized:
                doc.add_paragraph(f"• **Performance Optimization**: {len(underutilized)} running instances have very low CPU utilization (<5%). Consider downsizing or terminating unused instances.")
            
            if stopped_instances:
                doc.add_paragraph(f"• **Cost Optimization**: {len(stopped_instances)} instances are stopped but still incurring EBS storage costs. Consider terminating if no longer needed.")
            
            # Instance type recommendations
            old_generation_types = [inst for inst in instances_data 
                                  if any(inst.get('InstanceType', '').startswith(prefix) 
                                        for prefix in ['t1', 't2', 'm1', 'm2', 'm3', 'c1', 'c3'])]
            
            if old_generation_types:
                doc.add_paragraph(f"• **Instance Modernization**: {len(old_generation_types)} instances are using older generation instance types. Consider upgrading for better performance and cost efficiency.")
            
            # Security recommendations
            no_key_instances = [inst for inst in instances_data if inst.get('KeyName', 'None') == 'None']
            if no_key_instances:
                doc.add_paragraph(f"• **Key Pair Security**: {len(no_key_instances)} instances have no key pair assigned. Consider using AWS Systems Manager Session Manager for secure access.")
            
            untagged_instances = [inst for inst in instances_data if inst.get('TagCount', 0) == 0]
            if untagged_instances:
                doc.add_paragraph(f"• **Resource Tagging**: {len(untagged_instances)} instances have no tags. Implement consistent tagging for better resource management and cost allocation.")
            
            no_iam_role = [inst for inst in instances_data if inst.get('IamRole', 'None') == 'None']
            if no_iam_role:
                doc.add_paragraph(f"• **IAM Security**: {len(no_iam_role)} instances have no IAM role assigned. Consider using IAM roles instead of storing credentials on instances.")
            
            unencrypted_volumes = []
            for inst in instances_data:
                for volume in inst.get('ebs_volumes', []):
                    if not volume.get('encrypted', False):
                        unencrypted_volumes.append(inst)
                        break
            
            if unencrypted_volumes:
                doc.add_paragraph(f"• **Encryption Security**: {len(unencrypted_volumes)} instances have unencrypted EBS volumes. Enable encryption for data at rest protection.")
            
            monitoring_disabled = [inst for inst in instances_data if inst.get('MonitoringState', 'disabled') == 'disabled']
            if monitoring_disabled:
                doc.add_paragraph(f"• **Monitoring**: {len(monitoring_disabled)} instances have detailed monitoring disabled. Enable for better performance insights and alerting.")
            
            # Operating system recommendations
            old_os_instances = []
            for inst in instances_data:
                os = inst.get('OperatingSystem', '').lower()
                if any(old_version in os for old_version in ['server 2012', 'server 2016', 'ubuntu 18.04', 'centos 7', 'amazon linux 2']):
                    old_os_instances.append(inst)
            
            if old_os_instances:
                doc.add_paragraph(f"• **OS Updates**: {len(old_os_instances)} instances are running older operating system versions. Consider upgrading for security patches and performance improvements.")
            
            # Public IP recommendations
            public_ip_instances = [inst for inst in instances_data if inst.get('PublicIpAddress', 'None') != 'None']
            if public_ip_instances:
                doc.add_paragraph(f"• **Network Security**: {len(public_ip_instances)} instances have public IP addresses. Review if public access is necessary and ensure proper security group configurations.")
            
            # AMI recommendations
            public_ami_instances = []
            old_ami_instances = []
            
            for inst in instances_data:
                ami_details = inst.get('AmiDetails', {})
                if ami_details.get('Public', False) and ami_details.get('Owner', '') not in ['amazon', '099720109477']:  # Amazon and Canonical
                    public_ami_instances.append(inst)
                
                # Check AMI age (if creation date is available)
                creation_date = ami_details.get('CreationDate', '')
                if creation_date and creation_date != 'Unknown':
                    try:
                        # Simple date parsing for common AWS formats
                        if 'T' in creation_date:
                            ami_date = datetime.strptime(creation_date[:19], '%Y-%m-%dT%H:%M:%S')
                            if (datetime.now() - ami_date).days > 365:  # Older than 1 year
                                old_ami_instances.append(inst)
                    except:
                        pass
            
            if public_ami_instances:
                doc.add_paragraph(f"• **AMI Security**: {len(public_ami_instances)} instances are using public AMIs from third-party accounts. Consider using official AMIs or creating your own for better security control.")
            
            if old_ami_instances:
                doc.add_paragraph(f"• **AMI Currency**: {len(old_ami_instances)} instances are using AMIs older than 1 year. Consider updating to newer AMI versions for latest security patches.")

        # Save document
        filename = f"EC2_Comprehensive_Report_{self.account_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        doc.save(filename)
        print(f"\nComprehensive EC2 report saved successfully as: {filename}")

def main():
    parser = argparse.ArgumentParser(description='Generate Comprehensive EC2 Cost and Configuration Report.')
    parser.add_argument('--profile', type=str, default='default', help='AWS profile to use.')
    parser.add_argument('--region', type=str, default='us-east-1', help='Primary AWS region to use (default: us-east-1).')
    parser.add_argument('--start-month', type=int, default=4, help='Start month for cost analysis (1-12).')
    parser.add_argument('--end-month', type=int, default=6, help='End month for cost analysis (1-12).')
    parser.add_argument('--year', type=int, default=2025, help='Year for cost analysis.')
    args = parser.parse_args()

    print(f"Using AWS profile: {args.profile}")
    print(f"Using primary AWS region: {args.region}")
    print(f"Cost analysis period: {args.start_month}/{args.year} - {args.end_month}/{args.year}")
    
    report_generator = ComprehensiveEC2CostReportGenerator(
        profile_name=args.profile,
        region=args.region,
        start_month=args.start_month,
        end_month=args.end_month,
        year=args.year
    )
    
    instances_data = report_generator.fetch_all_ec2_instances()
    cost_data = report_generator.get_ec2_cost_history()
    report_generator.generate_report(instances_data, cost_data)

if __name__ == '__main__':
    main()