#!/usr/bin/env python3
"""
Enhanced S3 Cost and Configuration Report Generator
Fetches comprehensive S3 bucket configurations, their actual sizes, costs for a specified 3-month period,
and generates a detailed Word document report.
"""

import boto3
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import sys
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
import time

class EnhancedS3CostReportGenerator:
    def __init__(self, profile_name='yes4all', region='us-east-1', start_month=4, end_month=6, year=2025):
        self.profile_name = profile_name
        self.region = region
        try:
            # Create session with profile and region
            self.session = boto3.Session(profile_name=self.profile_name, region_name=self.region)
            
            # Test the session by getting account ID
            sts_client = self.session.client('sts')
            self.account_id = sts_client.get_caller_identity()['Account']
            
            # Initialize other clients
            self.s3 = self.session.client('s3')
            self.ce = self.session.client('ce', region_name='us-east-1')  # Cost Explorer is only available in us-east-1
            self.cloudwatch = self.session.client('cloudwatch')
            
            print(f"Successfully connected to AWS using profile '{self.profile_name}' in region '{self.region}'")
            print(f"Account ID: {self.account_id}")
            
        except Exception as e:
            print(f"Error creating AWS session: {e}")
            print(f"Please ensure the AWS profile '{self.profile_name}' is configured correctly.")
            print(f"You can also try:")
            print(f"  1. Set a default region: aws configure set region {self.region} --profile {self.profile_name}")
            print(f"  2. Use a different region with --region parameter")
            print(f"  3. Check your AWS credentials: aws sts get-caller-identity --profile {self.profile_name}")
            sys.exit(1)
            
        self.start_date = datetime(year, start_month, 1)
        self.end_date = datetime(year, end_month + 1, 1)

    def get_bucket_size_metrics(self, bucket_name, bucket_region=None):
        """Get bucket size and object count from CloudWatch metrics."""
        try:
            # Use region-specific CloudWatch client if bucket region is known
            if bucket_region and bucket_region != self.region:
                cw_client = self.session.client('cloudwatch', region_name=bucket_region)
            else:
                cw_client = self.cloudwatch
            
            # Get bucket size in bytes
            size_response = cw_client.get_metric_statistics(
                Namespace='AWS/S3',
                MetricName='BucketSizeBytes',
                Dimensions=[
                    {'Name': 'BucketName', 'Value': bucket_name},
                    {'Name': 'StorageType', 'Value': 'StandardStorage'}
                ],
                StartTime=datetime.utcnow() - timedelta(days=7),
                EndTime=datetime.utcnow(),
                Period=86400,  # Daily
                Statistics=['Average']
            )
            
            # Get number of objects
            count_response = cw_client.get_metric_statistics(
                Namespace='AWS/S3',
                MetricName='NumberOfObjects',
                Dimensions=[
                    {'Name': 'BucketName', 'Value': bucket_name},
                    {'Name': 'StorageType', 'Value': 'AllStorageTypes'}
                ],
                StartTime=datetime.utcnow() - timedelta(days=7),
                EndTime=datetime.utcnow(),
                Period=86400,
                Statistics=['Average']
            )
            
            size_bytes = 0
            object_count = 0
            
            if size_response['Datapoints']:
                size_bytes = max([dp['Average'] for dp in size_response['Datapoints']])
            
            if count_response['Datapoints']:
                object_count = max([dp['Average'] for dp in count_response['Datapoints']])
            
            return {
                'size_bytes': int(size_bytes),
                'size_gb': round(size_bytes / (1024**3), 2),
                'size_tb': round(size_bytes / (1024**4), 4),
                'object_count': int(object_count)
            }
        except Exception as e:
            print(f"    Could not get metrics for bucket {bucket_name}: {e}")
            return {
                'size_bytes': 0,
                'size_gb': 0,
                'size_tb': 0,
                'object_count': 0
            }

    def get_bucket_detailed_config(self, bucket_name):
        """Get comprehensive bucket configuration."""
        config = {'Name': bucket_name}
        
        try:
            # Get bucket region first
            bucket_location = self.s3.get_bucket_location(Bucket=bucket_name)
            bucket_region = bucket_location.get('LocationConstraint') or 'us-east-1'
            config['Region'] = bucket_region
            
            # Create region-specific S3 client if needed
            if bucket_region != self.region:
                regional_s3 = self.session.client('s3', region_name=bucket_region)
            else:
                regional_s3 = self.s3
                
        except Exception as e:
            config['Region'] = 'Unknown'
            regional_s3 = self.s3
            print(f"    Could not determine region for bucket {bucket_name}: {e}")
        
        # Versioning
        try:
            versioning = regional_s3.get_bucket_versioning(Bucket=bucket_name)
            config['Versioning'] = versioning.get('Status', 'Disabled')
            config['MfaDelete'] = versioning.get('MfaDelete', 'Disabled')
        except Exception:
            config['Versioning'] = 'Unknown'
            config['MfaDelete'] = 'Unknown'
        
        # Encryption
        try:
            encryption = regional_s3.get_bucket_encryption(Bucket=bucket_name)
            config['Encryption'] = 'Enabled'
            rules = encryption.get('ServerSideEncryptionConfiguration', {}).get('Rules', [])
            if rules:
                config['EncryptionType'] = rules[0].get('ApplyServerSideEncryptionByDefault', {}).get('SSEAlgorithm', 'Unknown')
            else:
                config['EncryptionType'] = 'Unknown'
        except Exception:
            config['Encryption'] = 'Disabled'
            config['EncryptionType'] = 'None'
        
        # Public Access Block
        try:
            pab = regional_s3.get_public_access_block(Bucket=bucket_name)
            pab_config = pab.get('PublicAccessBlockConfiguration', {})
            config['PublicAccessBlock'] = 'Enabled' if any(pab_config.values()) else 'Disabled'
            config['BlockPublicAcls'] = pab_config.get('BlockPublicAcls', False)
            config['IgnorePublicAcls'] = pab_config.get('IgnorePublicAcls', False)
            config['BlockPublicPolicy'] = pab_config.get('BlockPublicPolicy', False)
            config['RestrictPublicBuckets'] = pab_config.get('RestrictPublicBuckets', False)
        except Exception:
            config['PublicAccessBlock'] = 'Not Set'
            config['BlockPublicAcls'] = 'Unknown'
            config['IgnorePublicAcls'] = 'Unknown'
            config['BlockPublicPolicy'] = 'Unknown'
            config['RestrictPublicBuckets'] = 'Unknown'
        
        # Lifecycle Configuration
        try:
            lifecycle = regional_s3.get_bucket_lifecycle_configuration(Bucket=bucket_name)
            config['LifecycleRules'] = len(lifecycle.get('Rules', []))
            config['LifecycleEnabled'] = 'Yes' if config['LifecycleRules'] > 0 else 'No'
        except Exception:
            config['LifecycleRules'] = 0
            config['LifecycleEnabled'] = 'No'
        
        # CORS Configuration
        try:
            cors = regional_s3.get_bucket_cors(Bucket=bucket_name)
            config['CorsRules'] = len(cors.get('CORSRules', []))
            config['CorsEnabled'] = 'Yes' if config['CorsRules'] > 0 else 'No'
        except Exception:
            config['CorsRules'] = 0
            config['CorsEnabled'] = 'No'
        
        # Website Configuration
        try:
            website = regional_s3.get_bucket_website(Bucket=bucket_name)
            config['WebsiteEnabled'] = 'Yes'
            config['IndexDocument'] = website.get('IndexDocument', {}).get('Suffix', 'Not Set')
        except Exception:
            config['WebsiteEnabled'] = 'No'
            config['IndexDocument'] = 'N/A'
        
        # Logging Configuration
        try:
            logging = regional_s3.get_bucket_logging(Bucket=bucket_name)
            if logging.get('LoggingEnabled'):
                config['LoggingEnabled'] = 'Yes'
                config['LoggingTarget'] = logging['LoggingEnabled'].get('TargetBucket', 'Unknown')
            else:
                config['LoggingEnabled'] = 'No'
                config['LoggingTarget'] = 'N/A'
        except Exception:
            config['LoggingEnabled'] = 'Unknown'
            config['LoggingTarget'] = 'Unknown'
        
        # Notification Configuration
        try:
            notification = regional_s3.get_bucket_notification_configuration(Bucket=bucket_name)
            total_configs = (len(notification.get('TopicConfigurations', [])) + 
                           len(notification.get('QueueConfigurations', [])) + 
                           len(notification.get('LambdaConfigurations', [])))
            config['NotificationConfigs'] = total_configs
            config['NotificationEnabled'] = 'Yes' if total_configs > 0 else 'No'
        except Exception:
            config['NotificationConfigs'] = 0
            config['NotificationEnabled'] = 'No'
        
        # Replication Configuration
        try:
            replication = regional_s3.get_bucket_replication(Bucket=bucket_name)
            config['ReplicationRules'] = len(replication.get('ReplicationConfiguration', {}).get('Rules', []))
            config['ReplicationEnabled'] = 'Yes' if config['ReplicationRules'] > 0 else 'No'
        except Exception:
            config['ReplicationRules'] = 0
            config['ReplicationEnabled'] = 'No'
        
        # Object Lock Configuration
        try:
            object_lock = regional_s3.get_object_lock_configuration(Bucket=bucket_name)
            config['ObjectLockEnabled'] = object_lock.get('ObjectLockConfiguration', {}).get('ObjectLockEnabled', 'Disabled')
        except Exception:
            config['ObjectLockEnabled'] = 'Disabled'
        
        # Tagging
        try:
            tagging = regional_s3.get_bucket_tagging(Bucket=bucket_name)
            config['TagCount'] = len(tagging.get('TagSet', []))
            config['Tags'] = {tag['Key']: tag['Value'] for tag in tagging.get('TagSet', [])}
        except Exception:
            config['TagCount'] = 0
            config['Tags'] = {}
        
        return config

    def fetch_single_bucket_info(self, bucket):
        """Fetch complete information for a single bucket."""
        bucket_name = bucket['Name']
        print(f"  Processing bucket: {bucket_name}")
        
        try:
            # Get detailed configuration
            config = self.get_bucket_detailed_config(bucket_name)
            
            # Add creation date
            config['CreationDate'] = bucket['CreationDate'].strftime("%Y-%m-%d")
            
            # Get size metrics using the bucket's region
            bucket_region = config.get('Region', self.region)
            size_info = self.get_bucket_size_metrics(bucket_name, bucket_region)
            config.update(size_info)
            
            return config
        except Exception as e:
            print(f"    Error processing bucket {bucket_name}: {e}")
            return {
                'Name': bucket_name,
                'CreationDate': bucket['CreationDate'].strftime("%Y-%m-%d"),
                'Error': str(e)
            }

    def fetch_s3_bucket_configs(self):
        """Fetches comprehensive configuration for all S3 buckets with parallel processing."""
        print("Fetching comprehensive S3 bucket configurations...")
        buckets_config = []
        
        try:
            buckets = self.s3.list_buckets()['Buckets']
            print(f"Found {len(buckets)} buckets. Processing with parallel execution...")
            
            # Use ThreadPoolExecutor for parallel processing
            with ThreadPoolExecutor(max_workers=10) as executor:
                future_to_bucket = {executor.submit(self.fetch_single_bucket_info, bucket): bucket for bucket in buckets}
                
                for future in as_completed(future_to_bucket):
                    bucket_config = future.result()
                    if bucket_config:
                        buckets_config.append(bucket_config)
                        
        except Exception as e:
            print(f"Error fetching bucket list: {e}")
            
        return buckets_config

    def get_s3_cost_history(self):
        """Gets S3 cost data for the specified period."""
        print(f"Fetching S3 cost data from {self.start_date.strftime('%Y-%m')} to {self.end_date.strftime('%Y-%m')}...")
        try:
            response = self.ce.get_cost_and_usage(
                TimePeriod={
                    'Start': self.start_date.strftime('%Y-%m-%d'),
                    'End': self.end_date.strftime('%Y-%m-%d')
                },
                Granularity='MONTHLY',
                Filter={
                    'Dimensions': {
                        'Key': 'SERVICE',
                        'Values': ['Amazon Simple Storage Service']
                    }
                },
                Metrics=['UnblendedCost'],
                GroupBy=[{'Type': 'DIMENSION', 'Key': 'USAGE_TYPE'}]
            )
            return response['ResultsByTime']
        except Exception as e:
            print(f"Error fetching cost data: {e}")
            return None

    def format_size(self, size_bytes):
        """Format size in human readable format."""
        if size_bytes == 0:
            return "0 B"
        
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} PB"

    def generate_report(self, bucket_configs, cost_data):
        """Generates a comprehensive Word document report."""
        print("Generating comprehensive Word report...")
        doc = Document()
        doc.add_heading('AWS S3 Comprehensive Configuration and Cost Report', 0)

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"AWS Account ID: {self.account_id}")
        doc.add_paragraph(f"Profile: {self.profile_name}")
        doc.add_paragraph(f"Reporting Period: {self.start_date.strftime('%B %Y')} - {(self.end_date - timedelta(days=1)).strftime('%B %Y')}")
        
        if bucket_configs:
            total_buckets = len(bucket_configs)
            total_size = sum(bucket.get('size_bytes', 0) for bucket in bucket_configs)
            total_objects = sum(bucket.get('object_count', 0) for bucket in bucket_configs)
            encrypted_buckets = sum(1 for bucket in bucket_configs if bucket.get('Encryption') == 'Enabled')
            versioned_buckets = sum(1 for bucket in bucket_configs if bucket.get('Versioning') == 'Enabled')
            
            doc.add_paragraph(f"Total S3 Buckets: {total_buckets}")
            doc.add_paragraph(f"Total Storage Used: {self.format_size(total_size)}")
            doc.add_paragraph(f"Total Objects: {total_objects:,}")
            doc.add_paragraph(f"Encrypted Buckets: {encrypted_buckets} ({(encrypted_buckets/total_buckets*100):.1f}%)")
            doc.add_paragraph(f"Versioned Buckets: {versioned_buckets} ({(versioned_buckets/total_buckets*100):.1f}%)")

        # Bucket Overview Table
        doc.add_heading('S3 Bucket Overview', level=1)
        if bucket_configs:
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ['Bucket Name', 'Region', 'Size', 'Objects', 'Versioning', 'Encryption', 'Public Access', 'Creation Date']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            for bucket in sorted(bucket_configs, key=lambda x: x.get('size_bytes', 0), reverse=True):
                row_cells = table.add_row().cells
                row_cells[0].text = bucket.get('Name', 'Unknown')
                row_cells[1].text = bucket.get('Region', 'Unknown')
                row_cells[2].text = self.format_size(bucket.get('size_bytes', 0))
                row_cells[3].text = f"{bucket.get('object_count', 0):,}"
                row_cells[4].text = bucket.get('Versioning', 'Unknown')
                row_cells[5].text = bucket.get('Encryption', 'Unknown')
                row_cells[6].text = bucket.get('PublicAccessBlock', 'Unknown')
                row_cells[7].text = bucket.get('CreationDate', 'Unknown')

        # Detailed Configuration
        doc.add_heading('Detailed Bucket Configurations', level=1)
        for bucket in sorted(bucket_configs, key=lambda x: x.get('Name', '')):
            doc.add_heading(f"Bucket: {bucket.get('Name', 'Unknown')}", level=2)
            
            # Basic Info
            doc.add_heading('Basic Information', level=3)
            basic_info = [
                f"Region: {bucket.get('Region', 'Unknown')}",
                f"Creation Date: {bucket.get('CreationDate', 'Unknown')}",
                f"Size: {self.format_size(bucket.get('size_bytes', 0))}",
                f"Object Count: {bucket.get('object_count', 0):,}",
            ]
            for info in basic_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # Security Configuration
            doc.add_heading('Security Configuration', level=3)
            security_info = [
                f"Versioning: {bucket.get('Versioning', 'Unknown')}",
                f"MFA Delete: {bucket.get('MfaDelete', 'Unknown')}",
                f"Encryption: {bucket.get('Encryption', 'Unknown')}",
                f"Encryption Type: {bucket.get('EncryptionType', 'Unknown')}",
                f"Public Access Block: {bucket.get('PublicAccessBlock', 'Unknown')}",
                f"Object Lock: {bucket.get('ObjectLockEnabled', 'Unknown')}",
            ]
            for info in security_info:
                doc.add_paragraph(info, style='List Bullet')
            
            # Advanced Configuration
            doc.add_heading('Advanced Configuration', level=3)
            advanced_info = [
                f"Lifecycle Rules: {bucket.get('LifecycleRules', 0)}",
                f"CORS Rules: {bucket.get('CorsRules', 0)}",
                f"Website Hosting: {bucket.get('WebsiteEnabled', 'Unknown')}",
                f"Logging: {bucket.get('LoggingEnabled', 'Unknown')}",
                f"Notification Configs: {bucket.get('NotificationConfigs', 0)}",
                f"Replication Rules: {bucket.get('ReplicationRules', 0)}",
                f"Tags: {bucket.get('TagCount', 0)}",
            ]
            for info in advanced_info:
                doc.add_paragraph(info, style='List Bullet')

        # Cost Analysis
        doc.add_heading(f"S3 Cost Analysis ({self.start_date.strftime('%B')} - {(self.end_date - timedelta(days=1)).strftime('%B %Y')})", level=1)
        if cost_data:
            total_cost = 0
            for month_data in cost_data:
                month = datetime.strptime(month_data['TimePeriod']['Start'], '%Y-%m-%d').strftime('%B %Y')
                month_total = 0
                doc.add_heading(f"Costs for {month}", level=2)
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Usage Type'
                hdr_cells[1].text = 'Cost (USD)'
                hdr_cells[0].paragraphs[0].runs[0].font.bold = True
                hdr_cells[1].paragraphs[0].runs[0].font.bold = True

                groups = sorted(month_data['Groups'], key=lambda x: float(x['Metrics']['UnblendedCost']['Amount']), reverse=True)

                for group in groups:
                    cost = float(group['Metrics']['UnblendedCost']['Amount'])
                    if cost > 0:
                        row_cells = table.add_row().cells
                        row_cells[0].text = group['Keys'][0]
                        row_cells[1].text = f"${cost:,.2f}"
                        month_total += cost
                
                total_cost += month_total
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.add_run(f'Total cost for {month}: ').bold = True
                p.add_run(f"${month_total:,.2f}")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run('Total S3 cost for the period: ').bold = True
            p.add_run(f"${total_cost:,.2f}")
        else:
            doc.add_paragraph("Could not retrieve S3 cost data.")

        # Save document
        filename = f"S3_Comprehensive_Report_{self.account_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        doc.save(filename)
        print(f"\nComprehensive report saved successfully as: {filename}")

def main():
    parser = argparse.ArgumentParser(description='Generate Enhanced S3 Cost and Configuration Report.')
    parser.add_argument('--profile', type=str, default='yes4all', help='AWS profile to use.')
    parser.add_argument('--region', type=str, default='us-east-1', help='AWS region to use (default: us-east-1).')
    parser.add_argument('--start-month', type=int, default=4, help='Start month for cost analysis (1-12).')
    parser.add_argument('--end-month', type=int, default=6, help='End month for cost analysis (1-12).')
    parser.add_argument('--year', type=int, default=2025, help='Year for cost analysis.')
    args = parser.parse_args()

    print(f"Using AWS profile: {args.profile}")
    print(f"Using AWS region: {args.region}")
    print(f"Cost analysis period: {args.start_month}/{args.year} - {args.end_month}/{args.year}")
    
    report_generator = EnhancedS3CostReportGenerator(
        profile_name=args.profile,
        region=args.region,
        start_month=args.start_month,
        end_month=args.end_month,
        year=args.year
    )
    
    bucket_configs = report_generator.fetch_s3_bucket_configs()
    cost_data = report_generator.get_s3_cost_history()
    report_generator.generate_report(bucket_configs, cost_data)

if __name__ == '__main__':
    main()